from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np
from PIL import Image as PILImage
from pathlib import Path
import tempfile
import re

class DocumentGenerator:
    def __init__(self):
        self.temp_images = []
    
    def generate(self, ocr_result: dict, original_image_path: str, output_dir: Path) -> str:
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        for block in ocr_result['text_blocks']:
            text = block['text']
            block_type = block['type']
            level = block.get('level', 0)
            is_equation = block.get('is_equation', False)
            
            if is_equation:
                self._add_equation(doc, text)
            elif block_type == 'heading':
                self._add_heading(doc, text, level)
            elif block_type == 'list_item':
                self._add_list_item(doc, text)
            else:
                self._add_paragraph(doc, text)
        
        # CRITICAL: Diagrams are SKIPPED - no images should be inserted into documents
        # Only text extracted via OCR should be included
        # if 'diagrams' in ocr_result and ocr_result['diagrams']:
        #     for diag_idx, diagram in enumerate(ocr_result['diagrams']):
        #         self._add_diagram(doc, diagram['image'], f"Diagram {diag_idx + 1}")
        
        output_path = output_dir / f"output_{Path(original_image_path).stem}.docx"
        doc.save(str(output_path))
        
        return str(output_path)
    
    def _add_heading(self, doc: Document, text: str, level: int):
        text = re.sub(r'^#+\s*', '', text)
        text = re.sub(r'^[①②③④⑤⑥⑦⑧⑨⑩]\s*', '', text)
        heading = doc.add_heading(text, level=min(level, 3))
        
        for run in heading.runs:
            run.font.name = 'Calibri'
            if level == 1:
                run.font.size = Pt(16)
                run.font.bold = True
            elif level == 2:
                run.font.size = Pt(14)
                run.font.bold = True
            else:
                run.font.size = Pt(12)
                run.font.bold = True
    
    def _add_paragraph(self, doc: Document, text: str):
        text = self._clean_text(text)
        para = doc.add_paragraph(text)
        para.style = 'Normal'
    
    def _add_list_item(self, doc: Document, text: str):
        text = re.sub(r'^[•·◦▪▫]\s*', '', text)
        text = re.sub(r'^\(\w+\)\s*', '', text)
        text = re.sub(r'^\d+[\.\)]\s*', '', text)
        text = self._clean_text(text)
        para = doc.add_paragraph(text, style='List Bullet')
    
    def _add_equation(self, doc: Document, equation_text: str):
        equation_text = self._clean_text(equation_text)
        para = doc.add_paragraph()
        
        run = para.add_run(equation_text)
        run.font.italic = True
        run.font.name = 'Cambria Math'
        run.font.size = Pt(11)
        
        run2 = para.add_run(" (equation)")
        run2.font.italic = False
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_diagram(self, doc: Document, diagram_image: np.ndarray, caption: str):
        """
        DISABLED: No images should be inserted into documents.
        This method is kept for API compatibility but does nothing.
        Only text extracted via OCR should be included in documents.
        
        Args:
            doc: Word document (ignored)
            diagram_image: Diagram image array (ignored)
            caption: Caption text (ignored)
        """
        # CRITICAL: Images are NEVER inserted into documents
        # If OCR fails for a region, it is skipped entirely
        pass
    
    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = text.replace('|', 'l').replace('0', 'O')
        return text.strip()
